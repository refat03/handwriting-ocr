import streamlit as st
import easyocr
from docx import Document
from io import BytesIO
from PIL import Image
import numpy as np

# Page Config
st.set_page_config(page_title="Handwriting to Word", page_icon="📝")
st.title("📝 Handwritten to Word Converter")
st.write("Upload a photo of your notes and download them as a .docx file.")

# OCR Engine Loader (Cached so it's fast)
@st.cache_resource
def load_ocr():
    return easyocr.Reader(['en'])

reader = load_ocr()

# File Uploader
uploaded_file = st.file_uploader("Take a photo or upload an image", type=["jpg", "jpeg", "png"])

if uploaded_file:
    image = Image.open(uploaded_file)
    st.image(image, caption="Uploaded Image", use_container_width=True)

    if st.button("Generate Word Document"):
        with st.spinner("Processing text..."):
            # Run OCR
            img_array = np.array(image)
            result = reader.readtext(img_array, detail=0)
            full_text = "\n".join(result)

            # Create Word Doc
            doc = Document()
            doc.add_heading('Converted Handwritten Notes', 0)
            doc.add_paragraph(full_text)

            # Prepare for download
            doc_io = BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)

            st.success("Conversion Complete!")
            st.download_button(
                label="📥 Download Word File",
                data=doc_io,
                file_name="MyNotes.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
